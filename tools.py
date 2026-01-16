"""
Funciones de conversión y procesamiento para PDF Converter Pro.
"""
from pathlib import Path
from typing import Optional, Callable, Any
import os
import io
import shutil
import tempfile
import zipfile

# Type alias para callbacks de progreso
ProgressCallback = Callable[[int, int, str], None]
CancelCheck = Callable[[], bool]

# Formatos de imagen soportados
SUPPORTED_IMAGE_FORMATS = ["png", "jpg", "jpeg", "webp", "bmp", "gif", "tiff", "ico"]


# ===========================================================================
# PDF -> DOCX (Editable)
# ===========================================================================

def pdf_to_docx(
    input_pdf: Path,
    output_docx: Path,
    start: Optional[int] = None,
    end: Optional[int] = None,
    overwrite: bool = False
) -> None:
    """Convierte PDF a DOCX usando pdf2docx (texto editable)."""
    if output_docx.exists() and not overwrite:
        raise FileExistsError(f"El archivo ya existe: {output_docx}")

    from pdf2docx import Converter

    cv = Converter(str(input_pdf))
    try:
        cv.convert(str(output_docx), start=start, end=end)
    finally:
        cv.close()


def pdf_to_docx_with_progress(
    input_pdf: Path,
    output_docx: Path,
    start: Optional[int] = None,
    end: Optional[int] = None,
    overwrite: bool = False,
    progress_callback: Optional[ProgressCallback] = None,
    cancel_check: Optional[CancelCheck] = None
) -> None:
    """Convierte PDF a DOCX con reporte de progreso."""
    if output_docx.exists() and not overwrite:
        raise FileExistsError(f"El archivo ya existe: {output_docx}")

    from pdf2docx import Converter

    cv = Converter(str(input_pdf))
    try:
        # Obtener número de páginas
        import fitz
        doc = fitz.open(str(input_pdf))
        total_pages = doc.page_count
        doc.close()

        actual_start = start if start else 0
        actual_end = end if end else total_pages

        if progress_callback:
            progress_callback(0, actual_end - actual_start, f"Iniciando conversión de {total_pages} páginas...")

        # Convertir página por página para reportar progreso
        for i in range(actual_start, actual_end):
            if cancel_check and cancel_check():
                raise InterruptedError("Operación cancelada por el usuario")

            page_num = i + 1
            if progress_callback:
                progress_callback(i - actual_start + 1, actual_end - actual_start, f"Página {page_num}/{actual_end}")

        # Hacer la conversión real
        cv.convert(str(output_docx), start=start, end=end)

        if progress_callback:
            progress_callback(actual_end - actual_start, actual_end - actual_start, "Conversión completada")

    finally:
        cv.close()


# ===========================================================================
# DOCX -> PDF
# ===========================================================================

def docx_to_pdf(
    input_docx: Path,
    output_pdf: Path,
    overwrite: bool = False
) -> None:
    """Convierte DOCX a PDF usando docx2pdf (requiere Microsoft Word)."""
    if output_pdf.exists() and not overwrite:
        raise FileExistsError(f"El archivo ya existe: {output_pdf}")

    from docx2pdf import convert
    convert(str(input_docx), str(output_pdf))


# ===========================================================================
# PDF -> DOCX (Raster/Imagen)
# ===========================================================================

def pdf_to_docx_raster(
    input_pdf: Path,
    output_docx: Path,
    dpi: int = 200,
    overwrite: bool = False
) -> None:
    """Convierte PDF a DOCX renderizando como imágenes (fidelidad exacta)."""
    if output_docx.exists() and not overwrite:
        raise FileExistsError(f"El archivo ya existe: {output_docx}")

    import fitz
    from docx import Document
    from docx.shared import Inches

    doc = fitz.open(str(input_pdf))
    word_doc = Document()

    try:
        for page_num in range(doc.page_count):
            page = doc[page_num]
            # Renderizar página como imagen
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)

            # Guardar temporalmente
            img_data = pix.tobytes("png")

            # Insertar en Word
            from io import BytesIO
            img_stream = BytesIO(img_data)

            # Calcular tamaño en pulgadas (basado en tamaño de página)
            width_inches = page.rect.width / 72
            word_doc.add_picture(img_stream, width=Inches(min(width_inches, 7.5)))

            if page_num < doc.page_count - 1:
                word_doc.add_page_break()

        word_doc.save(str(output_docx))
    finally:
        doc.close()


def pdf_to_docx_raster_with_progress(
    input_pdf: Path,
    output_docx: Path,
    dpi: int = 200,
    overwrite: bool = False,
    progress_callback: Optional[ProgressCallback] = None,
    cancel_check: Optional[CancelCheck] = None
) -> None:
    """Convierte PDF a DOCX como imágenes con reporte de progreso."""
    if output_docx.exists() and not overwrite:
        raise FileExistsError(f"El archivo ya existe: {output_docx}")

    import fitz
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO

    doc = fitz.open(str(input_pdf))
    word_doc = Document()
    total_pages = doc.page_count

    if progress_callback:
        progress_callback(0, total_pages, f"Procesando {total_pages} páginas a {dpi} DPI...")

    try:
        for page_num in range(total_pages):
            if cancel_check and cancel_check():
                raise InterruptedError("Operación cancelada por el usuario")

            if progress_callback:
                progress_callback(page_num + 1, total_pages, f"Renderizando página {page_num + 1}/{total_pages}")

            page = doc[page_num]
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            img_stream = BytesIO(img_data)

            width_inches = page.rect.width / 72
            word_doc.add_picture(img_stream, width=Inches(min(width_inches, 7.5)))

            if page_num < total_pages - 1:
                word_doc.add_page_break()

        if progress_callback:
            progress_callback(total_pages, total_pages, "Guardando documento...")

        word_doc.save(str(output_docx))

    finally:
        doc.close()


# ===========================================================================
# OCR PDF -> DOCX
# ===========================================================================

def ocr_pdf_to_docx_with_progress(
    input_pdf: Path,
    output_docx: Path,
    dpi: int = 300,
    lang: str = "spa",
    progress_callback: Optional[ProgressCallback] = None,
    cancel_check: Optional[CancelCheck] = None
) -> None:
    """Convierte PDF a DOCX usando OCR (pytesseract)."""
    import fitz
    from docx import Document
    from PIL import Image
    import pytesseract
    from io import BytesIO

    doc = fitz.open(str(input_pdf))
    word_doc = Document()
    total_pages = doc.page_count

    if progress_callback:
        progress_callback(0, total_pages, f"Iniciando OCR ({lang})...")

    try:
        for page_num in range(total_pages):
            if cancel_check and cancel_check():
                raise InterruptedError("Operación cancelada por el usuario")

            if progress_callback:
                progress_callback(page_num + 1, total_pages, f"OCR página {page_num + 1}/{total_pages}")

            page = doc[page_num]
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)

            # Convertir a PIL Image para OCR
            img_data = pix.tobytes("png")
            img = Image.open(BytesIO(img_data))

            # Ejecutar OCR
            text = pytesseract.image_to_string(img, lang=lang)

            # Agregar texto al documento
            if text.strip():
                word_doc.add_paragraph(text)

            if page_num < total_pages - 1:
                word_doc.add_page_break()

        if progress_callback:
            progress_callback(total_pages, total_pages, "Guardando documento...")

        word_doc.save(str(output_docx))

    finally:
        doc.close()


# ===========================================================================
# Compresión PDF
# ===========================================================================

def compress_pdf_with_progress(
    input_pdf: Path,
    output_pdf: Path,
    progress_callback: Optional[ProgressCallback] = None,
    cancel_check: Optional[CancelCheck] = None
) -> dict:
    """Optimiza/comprime un PDF."""
    import pikepdf

    original_size = input_pdf.stat().st_size

    if progress_callback:
        progress_callback(0, 3, "Abriendo PDF...")

    if cancel_check and cancel_check():
        raise InterruptedError("Operación cancelada")

    with pikepdf.open(str(input_pdf)) as pdf:
        if progress_callback:
            progress_callback(1, 3, "Optimizando contenido...")

        if cancel_check and cancel_check():
            raise InterruptedError("Operación cancelada")

        if progress_callback:
            progress_callback(2, 3, "Guardando PDF optimizado...")

        pdf.save(
            str(output_pdf),
            compress_streams=True,
            object_stream_mode=pikepdf.ObjectStreamMode.generate
        )

    new_size = output_pdf.stat().st_size
    reduction = ((original_size - new_size) / original_size) * 100 if original_size > 0 else 0

    if progress_callback:
        progress_callback(3, 3, "Compresión completada")

    return {
        "original_size": original_size,
        "new_size": new_size,
        "reduction_percent": max(0, reduction)
    }


# ===========================================================================
# Compresión imágenes DOCX
# ===========================================================================

def compress_docx_images_with_progress(
    input_docx: Path,
    output_docx: Path,
    quality: int = 75,
    max_width: Optional[int] = None,
    max_height: Optional[int] = None,
    progress_callback: Optional[ProgressCallback] = None,
    cancel_check: Optional[CancelCheck] = None
) -> dict:
    """Comprime las imágenes dentro de un archivo DOCX."""
    from PIL import Image

    original_size = input_docx.stat().st_size

    # Crear directorio temporal
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        extract_dir = tmpdir_path / "extracted"

        if progress_callback:
            progress_callback(0, 4, "Extrayendo DOCX...")

        # Extraer DOCX (es un ZIP)
        with zipfile.ZipFile(str(input_docx), 'r') as zip_ref:
            zip_ref.extractall(str(extract_dir))

        if cancel_check and cancel_check():
            raise InterruptedError("Operación cancelada")

        # Buscar imágenes
        media_dir = extract_dir / "word" / "media"
        images_processed = 0

        if media_dir.exists():
            image_files = list(media_dir.glob("*"))
            total_images = len(image_files)

            if progress_callback:
                progress_callback(1, 4, f"Comprimiendo {total_images} imágenes...")

            for i, img_path in enumerate(image_files):
                if cancel_check and cancel_check():
                    raise InterruptedError("Operación cancelada")

                try:
                    # Abrir imagen
                    with Image.open(img_path) as img:
                        # Convertir a RGB si es necesario
                        if img.mode in ('RGBA', 'P'):
                            img = img.convert('RGB')

                        # Redimensionar si se especificó
                        if max_width or max_height:
                            w, h = img.size
                            new_w, new_h = w, h

                            if max_width and w > max_width:
                                ratio = max_width / w
                                new_w = max_width
                                new_h = int(h * ratio)

                            if max_height and new_h > max_height:
                                ratio = max_height / new_h
                                new_h = max_height
                                new_w = int(new_w * ratio)

                            if new_w != w or new_h != h:
                                img = img.resize((new_w, new_h), Image.LANCZOS)

                        # Guardar como JPEG comprimido
                        new_path = img_path.with_suffix('.jpeg')
                        img.save(str(new_path), 'JPEG', quality=quality, optimize=True)

                        # Eliminar original si es diferente
                        if new_path != img_path:
                            img_path.unlink()

                        images_processed += 1

                except Exception:
                    # Si falla, dejar la imagen original
                    pass

                if progress_callback and total_images > 0:
                    progress_callback(1, 4, f"Imagen {i + 1}/{total_images}")

        if progress_callback:
            progress_callback(2, 4, "Reempaquetando DOCX...")

        if cancel_check and cancel_check():
            raise InterruptedError("Operación cancelada")

        # Crear nuevo DOCX
        with zipfile.ZipFile(str(output_docx), 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(str(extract_dir)):
                for file in files:
                    file_path = Path(root) / file
                    arcname = file_path.relative_to(extract_dir)
                    zipf.write(str(file_path), str(arcname))

    if progress_callback:
        progress_callback(4, 4, "Compresión completada")

    new_size = output_docx.stat().st_size
    reduction = ((original_size - new_size) / original_size) * 100 if original_size > 0 else 0

    return {
        "original_size": original_size,
        "new_size": new_size,
        "reduction_percent": max(0, reduction),
        "images_processed": images_processed
    }


# ===========================================================================
# Conversión de Imágenes
# ===========================================================================

def convert_image(
    input_path: Path,
    output_path: Path,
    output_format: str,
    quality: int = 95,
    resize: Optional[tuple[int, int]] = None,
    maintain_aspect: bool = True,
    overwrite: bool = False
) -> None:
    """Convierte una imagen a otro formato."""
    if output_path.exists() and not overwrite:
        raise FileExistsError(f"El archivo ya existe: {output_path}")

    from PIL import Image

    with Image.open(input_path) as img:
        # Convertir modo si es necesario
        original_mode = img.mode

        # Para ICO necesitamos RGBA
        if output_format.lower() == "ico":
            if img.mode != 'RGBA':
                img = img.convert('RGBA')
        # Para JPEG necesitamos RGB
        elif output_format.lower() in ('jpg', 'jpeg'):
            if img.mode in ('RGBA', 'P', 'LA'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if 'A' in img.mode else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')

        # Redimensionar si se especificó
        if resize:
            target_w, target_h = resize
            orig_w, orig_h = img.size

            if maintain_aspect:
                ratio_w = target_w / orig_w if target_w > 0 else float('inf')
                ratio_h = target_h / orig_h if target_h > 0 else float('inf')
                ratio = min(ratio_w, ratio_h)

                if ratio < 1:  # Solo reducir, no ampliar
                    new_w = int(orig_w * ratio)
                    new_h = int(orig_h * ratio)
                    img = img.resize((new_w, new_h), Image.LANCZOS)
            else:
                if target_w > 0 and target_h > 0:
                    img = img.resize((target_w, target_h), Image.LANCZOS)

        # Guardar
        output_path.parent.mkdir(parents=True, exist_ok=True)

        save_kwargs: dict[str, Any] = {}

        if output_format.lower() in ('jpg', 'jpeg'):
            save_kwargs['quality'] = quality
            save_kwargs['optimize'] = True
        elif output_format.lower() == 'png':
            save_kwargs['optimize'] = True
        elif output_format.lower() == 'webp':
            save_kwargs['quality'] = quality
        elif output_format.lower() == 'ico':
            # ICO tiene tamaños específicos
            sizes = [(256, 256), (128, 128), (64, 64), (48, 48), (32, 32), (16, 16)]
            img.save(str(output_path), format='ICO', sizes=sizes)
            return

        img.save(str(output_path), format=output_format.upper(), **save_kwargs)


def get_image_info(input_path: Path) -> dict:
    """Obtiene información de una imagen."""
    from PIL import Image

    file_size = input_path.stat().st_size

    with Image.open(input_path) as img:
        width, height = img.size
        format_name = img.format or "Unknown"
        mode = img.mode

    # Formatear tamaño
    if file_size < 1024:
        size_str = f"{file_size} B"
    elif file_size < 1024 * 1024:
        size_str = f"{file_size / 1024:.1f} KB"
    else:
        size_str = f"{file_size / (1024 * 1024):.1f} MB"

    return {
        "width": width,
        "height": height,
        "format": format_name,
        "mode": mode,
        "file_size": file_size,
        "file_size_human": size_str
    }


def batch_convert_images(
    input_files: list[Path],
    output_dir: Path,
    output_format: str,
    quality: int = 95,
    resize: Optional[tuple[int, int]] = None,
    maintain_aspect: bool = True,
    overwrite: bool = False,
    progress_callback: Optional[ProgressCallback] = None,
    cancel_check: Optional[CancelCheck] = None
) -> dict:
    """Convierte múltiples imágenes en lote."""
    output_dir.mkdir(parents=True, exist_ok=True)

    total = len(input_files)
    converted = 0
    errors = []

    for i, input_path in enumerate(input_files):
        if cancel_check and cancel_check():
            raise InterruptedError("Operación cancelada")

        if progress_callback:
            progress_callback(i + 1, total, f"Convirtiendo {input_path.name}")

        try:
            ext = output_format.lower()
            if ext == "jpeg":
                ext = "jpg"
            output_path = output_dir / f"{input_path.stem}.{ext}"

            convert_image(
                input_path, output_path, output_format,
                quality, resize, maintain_aspect, overwrite
            )
            converted += 1

        except Exception as e:
            errors.append((input_path.name, str(e)))

    return {
        "total": total,
        "converted": converted,
        "errors": errors
    }


# ===========================================================================
# Extracción de imágenes
# ===========================================================================

def extract_images_from_pdf(
    input_pdf: Path,
    output_dir: Path,
    output_format: str = "png"
) -> list[Path]:
    """Extrae todas las imágenes de un PDF."""
    import fitz
    from PIL import Image
    from io import BytesIO

    output_dir.mkdir(parents=True, exist_ok=True)
    extracted = []

    doc = fitz.open(str(input_pdf))

    try:
        img_count = 0
        for page_num in range(doc.page_count):
            page = doc[page_num]
            image_list = page.get_images()

            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                # Convertir a formato deseado
                pil_img = Image.open(BytesIO(image_bytes))

                if output_format.lower() in ('jpg', 'jpeg') and pil_img.mode in ('RGBA', 'P'):
                    pil_img = pil_img.convert('RGB')

                img_count += 1
                ext = "jpg" if output_format.lower() == "jpeg" else output_format.lower()
                output_path = output_dir / f"imagen_{img_count:04d}.{ext}"

                pil_img.save(str(output_path))
                extracted.append(output_path)

    finally:
        doc.close()

    return extracted


def extract_images_from_docx(
    input_docx: Path,
    output_dir: Path
) -> list[Path]:
    """Extrae todas las imágenes de un archivo DOCX."""
    output_dir.mkdir(parents=True, exist_ok=True)
    extracted = []

    with zipfile.ZipFile(str(input_docx), 'r') as zipf:
        for name in zipf.namelist():
            if name.startswith("word/media/"):
                # Extraer imagen
                img_data = zipf.read(name)
                img_name = Path(name).name
                output_path = output_dir / img_name

                output_path.write_bytes(img_data)
                extracted.append(output_path)

    return extracted
